import re
import matplotlib.pyplot as plt
import numpy as np
import requests
import sys
from PIL import Image
import docx
from docx.shared import RGBColor
from docx.shared import Inches


# task 2
def extractions():
    title_regex = re.compile(r'^Title: (.+)$')
    author_regex = re.compile(r'^Author: (.+)$')
    try:
        with open('Book.txt', 'r', encoding='utf-8') as f:
            lines = f.readlines()
            chapter = ''.join(lines[199:538])

            for line in lines:
                if title_regex.match(line):
                    title = title_regex.match(line).group(1)
                elif author_regex.match(line):
                    author = author_regex.match(line).group(1)
    except FileNotFoundError:
        print(f'The book file does not exist!')
        sys.exit()

    return title, author, chapter


def count_words(text: str):
    paragraphs = text.split('\n\n')
    lengths = []
    total_words = 0

    for paragraph in paragraphs:
        words = re.findall(r'\b\w+\b', paragraph)
        num_words = len(words)
        lengths.append(num_words)
        total_words += num_words

    return lengths, total_words


# task 3
def create_plot(lengths_list):
    lengths_max = np.max(lengths_list)
    lengths_min = np.min(lengths_list)
    lengths_mean = np.mean(lengths_list)

    x = np.arange(lengths_min, lengths_max + 1)

    y = []
    for i in x:
        y.append(lengths_list.count(i))

    fig = plt.figure(figsize=(20, 10))

    plt.plot(x, y, color='#2CF560')
    plt.xlabel('Number of words in paragraphs from min to max', fontsize=24)
    plt.ylabel('Number of paragraphs with given length', fontsize=24)
    plt.title('Distribution of Lengths of Paragraphs', fontsize=36)
    # I don't label every x value but only those for which number of paragraphs is > 0,
    # because otherwise the numbers will overlap producing unreadable picture
    plt.xticks(lengths_list, rotation=-45)
    plt.yticks(y)
    fig.savefig('plot.jpg')

    return lengths_mean, lengths_max, lengths_min


# task 4 and 5
def process_picture():
    url = 'https://www.gutenberg.org/cache/epub/6630/images/image01.jpg'
    response = requests.get(url)

    with open('picture1.jpg', 'wb') as f:
        f.write(response.content)

    image = Image.open(f.name)
    image = image.crop((10, 150, 640, 886))
    w, h = image.size
    image = image.resize((w+10, h+10))
    image.save('picture1_cropped.jpg')


# task 6
def process_logo():
    url = 'https://brand.esa.int/files/2020/05/ESA_logo_2020_Black-1024x643.jpg'
    response = requests.get(url)

    with open('logo.jpg', 'wb') as f:
        f.write(response.content)

    try:
        image = Image.open(f.name)
        image = image.crop((150, 150, 874, 493))
        w, h = image.size
        image = image.resize((int(w/5), int(h/5)))
        rotated = image.rotate(-90, expand=True)

        dest = Image.open('picture1_cropped.jpg')
        dest.paste(rotated, (50, 550))
        dest.save('picture1_final.jpg')
    except FileNotFoundError:
        print(f'The picture file does not exist!')
        sys.exit()


# task 7
def create_word_doc(title: str, author: str, plot_info: list):
    doc = docx.Document()
    doc.add_paragraph(title, 'Title')
    doc.add_paragraph(author)
    doc.paragraphs[1].runs[0].bold = True

    doc.paragraphs[1].runs[0].font.color.rgb = RGBColor(0, 0, 139)

    doc.add_picture('picture1_final.jpg', width=Inches(6.0))
    doc.add_paragraph('Report created by: Valentina Bolbas.')

    doc.paragraphs[3].runs[0].bold = True
    doc.paragraphs[3].runs[0].italic = True
    doc.paragraphs[3].runs[0].font.color.rgb = RGBColor(160, 32, 240)

    doc.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
    doc.add_picture('plot.jpg', width=Inches(6.0))
    paragraph = doc.add_paragraph()
    paragraph.add_run('This plot shows the ')

    paragraph.add_run('distribution of number of words')
    paragraph.runs[1].underline = True
    paragraph.runs[1].font.color.rgb = RGBColor(34, 201, 107)

    paragraph.add_run(' in each paragraph of the first chapter of the book. Y values show how many paragraphs are of the given length, which is represented by X values.')
    paragraph.add_run('\nTotal number of paragraphs is ')
    paragraph.add_run('20')
    paragraph.runs[4].bold = True
    paragraph.runs[4].font.color.rgb = RGBColor(185, 53, 212)

    paragraph.add_run(';\ntotal number of words in the first chapter is ')
    paragraph.add_run(f'{plot_info[0]}')
    paragraph.runs[6].bold = True
    paragraph.runs[6].font.color.rgb = RGBColor(49, 165, 173)

    paragraph.add_run(';\naverage (mean) number of words per paragraph is ')
    paragraph.add_run(f'{plot_info[1]}')
    paragraph.runs[8].bold = True
    paragraph.runs[8].font.color.rgb = RGBColor(165, 179, 59)

    paragraph.add_run(';\nmaximum number of words in paragraph is ')
    paragraph.add_run(f'{plot_info[2]}')
    paragraph.runs[10].bold = True
    paragraph.runs[10].font.color.rgb = RGBColor(176, 53, 84)

    paragraph.add_run(';\nminimum number of words in paragraph is ')
    paragraph.add_run(f'{plot_info[3]}')
    paragraph.runs[12].bold = True
    paragraph.runs[12].font.color.rgb = RGBColor(240, 137, 163)

    paragraph.add_run('.\nThe book was downloaded from: ')
    paragraph.add_run('https://www.gutenberg.org/ebooks/6630')
    paragraph.runs[14].underline = True
    paragraph.runs[14].font.color.rgb = RGBColor(62, 69, 189)

    doc.save('Report.docx')


def run():
    info_tuple = extractions()
    word_numbers = count_words(info_tuple[2])
    y_values = create_plot(word_numbers[0])
    process_picture()
    process_logo()
    plot_info = [word_numbers[1], y_values[0], y_values[1], y_values[2]]
    create_word_doc(info_tuple[0], info_tuple[1], plot_info)


if __name__ == "__main__":
    run()
